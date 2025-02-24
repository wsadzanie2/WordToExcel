from os.path import isfile
import os
from docx.document import Document as Doc_type
from docx.api import Document
import xlsxwriter


try:
    from pdf2docx import Converter
except Exception as e:
    print("wsparcie dla plików .pdf wyłączone. Brakuje biblioteki pdf2docx")
    print("Błąd programu: ")
    print(e)

if os.name == 'nt':
    PATH_SEP = "\\"
else:
    PATH_SEP = "/"
input_file = ''
output_file = ''

def get_spaces_between_words(line):
    counting_spaces = False
    for index, char in enumerate(line):
        if counting_spaces:
            if char != ' ':
                return index
        elif char == ' ':
            counting_spaces = True
                            
    return 0

try:
    import easygui # type: ignore
except Exception as e:
    print("Brakuje biblioteki easygui")
    print("na Windowsie trzeba wpisać pip install easygui, a na Ubuntu sudo apt-get install python3-easygui")
    raise e
loop = True
while loop:
    input_file = easygui.fileopenbox(filetypes=["*.docx", "*.pdf"], default="*.docx", title="Wybierz plik (.docx lub .pdf)") # type: ignore
    if input_file is None:
        exit()


    output_file = easygui.filesavebox(filetypes=["*.xlsx"], default=f"{input_file.split('.')[0]}.xlsx", title="Wybierz plik Worda (.xlsx)") # type: ignore
    if output_file is None:
        exit()
    if input_file.endswith(".pdf"): # type: ignore
        box = easygui.ccbox(msg=f"{input_file.split(PATH_SEP)[-1]} -> {input_file.split(PATH_SEP)[-1][:-4] + '.docx'} -> {output_file.split(PATH_SEP)[-1]}", choices=('Dobrze jest', 'Jeszcze raz')) # type: ignore
    else:
        box = easygui.ccbox(msg=f"{input_file.split(PATH_SEP)[-1]} -> {output_file.split(PATH_SEP)[-1]}", choices=('Dobrze jest', 'Jeszcze raz')) # type: ignore
    if box:
        loop = False


is_pdf = input_file.endswith(".pdf") # type: ignore



workbook = xlsxwriter.Workbook(output_file)
worksheet = workbook.add_worksheet()

if is_pdf:
    if isfile(input_file[:-4] + ".docx"): # type: ignore
        box = easygui.ccbox(msg=f"{input_file[:-4] + '.docx'} już istnieje. Zamiana pliku .pdf na .xlsx wyczyści ten plik.", choices=('Dobrze jest', 'ups...')) # type: ignore
        if not box:
            exit()
    cv = Converter(input_file)
    cv.convert(input_file[:-4] + ".docx") # type: ignore
    cv.close()
    # print(input_file[:-4] + ".docx") # type: ignore
    input_file = input_file[:-4] + ".docx" # type: ignore

print("Started working")
print(f"{input_file} -> {output_file}")

doc: Doc_type = Document(input_file) # type: ignore
tables = doc.tables

x = 0
y = 0

def getText(document):
    doc = document
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
        # fullText.append(para)
    # return fullText
    return '\n'.join(fullText)

for thing in doc.iter_inner_content():
    try:
        thing.runs # type: ignore
        x = 0
        if is_pdf:
            for index, line in enumerate(thing.text.split("\n")): # type: ignore
                if len(line) < 1 and index < 1:
                    print(index)
                    continue
                worksheet.write(y, x, line) # type: ignore
                y += 1
                print(line + " :)") # type: ignore
        else:
            worksheet.write(y, x, thing.text) # type: ignore
            print(thing.text + " :)") # type: ignore
            y += 1
    except AttributeError as e:
        # handle tables
        for row in thing.rows: # type: ignore
            for cell in row.cells:
                # print(cell.text, end='|')
                worksheet.write(y, x, cell.text)
                x += 1
            y += 1
            x = 0
            # print()

workbook.close()
